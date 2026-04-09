from __future__ import annotations

import csv
import io
from pathlib import Path

from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document
import xlrd

SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".xls", ".xlsx", ".doc", ".docx", ".pdf"}


def extract_text_from_source(path: str | Path, *, mime_type: str | None = None) -> str:
    source_path = Path(path)
    suffix = source_path.suffix.lower()
    if suffix not in SUPPORTED_TEXT_EXTENSIONS:
        raise ValueError(f"不支持的知识文件类型: {suffix or source_path.name}")

    if suffix in {".txt", ".md"}:
        return _read_text_with_fallbacks(source_path)
    if suffix == ".csv":
        return _extract_csv_text(source_path)
    if suffix == ".docx":
        return _extract_docx_text(source_path)
    if suffix == ".xlsx":
        return _extract_xlsx_text(source_path)
    if suffix == ".xls":
        return _extract_xls_text(source_path)
    if suffix == ".pdf":
        return _extract_pdf_text(source_path)
    if suffix == ".doc":
        raise ValueError("DOC 旧版二进制文档暂不支持直接解析，请转换为 DOCX / TXT / PDF。")
    raise ValueError(f"不支持的知识文件类型: {mime_type or suffix}")


def _read_text_with_fallbacks(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding).strip()
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def _extract_csv_text(path: Path) -> str:
    raw_text = _read_text_with_fallbacks(path)
    csv_buffer = io.StringIO(raw_text)
    reader = csv.DictReader(csv_buffer)
    rows = list(reader)
    if not rows:
        raise ValueError("CSV 文件为空，无法提取文本。")

    lines = [f"表格文件《{path.stem}》包含以下记录："]
    for index, row in enumerate(rows, start=1):
        summary = "；".join(f"{key}={value}" for key, value in row.items())
        lines.append(f"第 {index} 行：{summary}")
    return "\n".join(lines)


def _extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append("；".join(cells))
        if rows:
            tables.append(f"表 {table_index}：\n" + "\n".join(rows))
    text = "\n\n".join(paragraphs + tables).strip()
    if not text:
        raise ValueError("DOCX 文档没有可提取的文本。")
    return text


def _extract_xlsx_text(path: Path) -> str:
    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(cell).strip() if cell is not None else "" for cell in rows[0]]
        lines.append(f"工作表《{sheet.title}》")
        for index, row in enumerate(rows[1:] if headers else rows, start=1):
            values = []
            for header, cell in zip(headers, row, strict=False):
                if header and cell not in (None, ""):
                    values.append(f"{header}={cell}")
            if values:
                lines.append(f"第 {index} 行：{'；'.join(values)}")
    text = "\n".join(lines).strip()
    if not text:
        raise ValueError("XLSX 文件没有可提取的文本。")
    return text


def _extract_xls_text(path: Path) -> str:
    workbook = xlrd.open_workbook(str(path))
    lines: list[str] = []
    for sheet in workbook.sheets():
        if sheet.nrows == 0:
            continue
        headers = [str(sheet.cell_value(0, col)).strip() for col in range(sheet.ncols)]
        lines.append(f"工作表《{sheet.name}》")
        start_row = 1 if headers else 0
        for row_index in range(start_row, sheet.nrows):
            values = []
            for col_index in range(sheet.ncols):
                cell_value = sheet.cell_value(row_index, col_index)
                header = headers[col_index] if col_index < len(headers) else ""
                if header and cell_value not in ("", None):
                    values.append(f"{header}={cell_value}")
            if values:
                lines.append(f"第 {row_index + 1} 行：{'；'.join(values)}")
    text = "\n".join(lines).strip()
    if not text:
        raise ValueError("XLS 文件没有可提取的文本。")
    return text


def _extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        page_text = page_text.strip()
        if page_text:
            pages.append(page_text)
    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError("PDF 文件没有可提取的文本，可能是扫描件或空白文档。")
    return text
