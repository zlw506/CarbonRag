from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt

from app.report.export.schemas import ExportResult, ReportBlock, ReportDocumentIR, ReportTable
from app.report.export.security import sha256_file


class DocxReportExporter:
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def export(self, ir: ReportDocumentIR, output_path: Path, template_id: str | None = None) -> ExportResult:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        _configure_styles(document)

        document.add_heading(ir.title, level=0)
        if ir.subtitle:
            document.add_paragraph(ir.subtitle, style="Subtitle")
        if ir.metadata:
            _add_metadata(document, ir.metadata)

        for section in ir.sections:
            document.add_heading(section.heading, level=1)
            for block in section.blocks:
                _add_block(document, block)

        if ir.references:
            document.add_heading("引用依据", level=1)
            for index, citation in enumerate(ir.references, start=1):
                paragraph = document.add_paragraph(style="List Number")
                paragraph.add_run(f"{citation.title}").bold = True
                paragraph.add_run(f"｜{citation.source}")
                if citation.source_url:
                    paragraph.add_run(f"｜{citation.source_url}")
                if citation.snippet:
                    document.add_paragraph(citation.snippet, style="Quote")

        for appendix in ir.appendices:
            document.add_page_break()
            document.add_heading(appendix.heading, level=1)
            for block in appendix.blocks:
                _add_block(document, block)

        document.save(output_path)
        return ExportResult(
            output_path=output_path,
            content_type=self.content_type,
            file_size_bytes=output_path.stat().st_size,
            checksum_sha256=sha256_file(output_path),
        )


def _configure_styles(document: Document) -> None:
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        if style_name == "Normal":
            style.font.size = Pt(10.5)


def _add_metadata(document: Document, metadata: dict[str, str]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in metadata.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value
    document.add_paragraph()


def _add_block(document: Document, block: ReportBlock) -> None:
    if block.type == "paragraph":
        document.add_paragraph(block.text or "")
    elif block.type == "heading":
        document.add_heading(block.text or "", level=2)
    elif block.type == "bullet_list":
        for item in block.items or []:
            document.add_paragraph(item, style="List Bullet")
    elif block.type == "numbered_list":
        for item in block.items or []:
            document.add_paragraph(item, style="List Number")
    elif block.type == "quote":
        document.add_paragraph(block.text or "", style="Quote")
    elif block.type == "table" and block.table is not None:
        _add_table(document, block.table)
    elif block.type == "page_break":
        run = document.add_paragraph().add_run()
        run.add_break(WD_BREAK.PAGE)


def _add_table(document: Document, table_ir: ReportTable) -> None:
    if table_ir.caption:
        document.add_paragraph(table_ir.caption).runs[0].bold = True
    table = document.add_table(rows=1, cols=max(1, len(table_ir.columns)))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, column in enumerate(table_ir.columns):
        header_cells[index].text = column
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_values in table_ir.rows:
        row = table.add_row().cells
        for index, value in enumerate(row_values[: len(table_ir.columns)]):
            row[index].text = value
    if table_ir.footnote:
        document.add_paragraph(table_ir.footnote)
    document.add_paragraph()
