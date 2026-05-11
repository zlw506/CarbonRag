from pathlib import Path

import pytest
from docx import Document
from reportlab.pdfgen import canvas

from app.files.parser.registry import FileParserRegistry
from app.knowledge.parsers import KnowledgeParseError


class _UnavailableDocling:
    def supports(self, *args, **kwargs) -> bool:  # noqa: ANN002, ANN003
        return False

    def parse(self, *args, **kwargs):  # noqa: ANN002, ANN003
        raise AssertionError("fallback parser should be used")


def test_parser_registry_uses_lightweight_fallback_for_text(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.txt"
    file_path.write_text("企业电费账单显示 2026 年 1 月用电 1200 kWh。", encoding="utf-8")

    parsed = FileParserRegistry(docling_provider=_UnavailableDocling()).parse(
        path=file_path,
        mime_type="text/plain",
    )

    assert parsed.parser_name == "carbonrag-fallback"
    assert "1200 kWh" in parsed.text
    assert parsed.summary


def test_parser_registry_strips_html_with_fallback(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.html"
    file_path.write_text("<html><body><h1>排放报告</h1><p>用电量 800 kWh</p></body></html>", encoding="utf-8")

    parsed = FileParserRegistry(docling_provider=_UnavailableDocling()).parse(
        path=file_path,
        mime_type="text/html",
    )

    assert parsed.parser_name == "carbonrag-html-fallback"
    assert "排放报告" in parsed.text
    assert "<h1>" not in parsed.text


def test_parser_registry_parses_docx_with_fallback(tmp_path: Path) -> None:
    file_path = tmp_path / "carbon-report.docx"
    document = Document()
    document.add_paragraph("企业碳排放报告摘要")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "外购电力"
    table.rows[0].cells[1].text = "1200 kWh"
    document.save(file_path)

    parsed = FileParserRegistry(docling_provider=_UnavailableDocling()).parse(
        path=file_path,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert parsed.parser_name == "carbonrag-fallback"
    assert "企业碳排放报告摘要" in parsed.text
    assert "外购电力" in parsed.text
    assert "1200 kWh" in parsed.text


def test_parser_registry_parses_text_pdf_with_fallback(tmp_path: Path) -> None:
    file_path = tmp_path / "carbon-report.pdf"
    pdf = canvas.Canvas(str(file_path))
    pdf.drawString(72, 720, "Purchased electricity 1200 kWh")
    pdf.showPage()
    pdf.drawString(72, 720, "Natural gas 300 m3")
    pdf.save()

    parsed = FileParserRegistry(docling_provider=_UnavailableDocling()).parse(
        path=file_path,
        mime_type="application/pdf",
    )

    assert parsed.parser_name == "carbonrag-fallback"
    assert parsed.page_count == 2
    assert "Purchased electricity 1200 kWh" in parsed.text
    assert "Natural gas 300 m3" in parsed.text


def test_parser_registry_marks_image_as_ocr_unavailable_without_docling(tmp_path: Path) -> None:
    file_path = tmp_path / "scan.png"
    file_path.write_bytes(b"\x89PNG\r\n\x1a\n")

    with pytest.raises(KnowledgeParseError, match="Docling OCR"):
        FileParserRegistry(docling_provider=_UnavailableDocling()).parse(
            path=file_path,
            mime_type="image/png",
        )
