from datetime import datetime, timezone

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.report.export.docx_exporter import DocxReportExporter
from app.report.export.markdown_ir import report_to_ir
from app.report.export.pdf_exporter import PdfReportExporter
from app.report.export.schemas import CreateReportExportRequest
from app.report.export.service import ReportExportService
from app.report.export.storage import ReportExportStorage
from app.report.schemas import ReportCitation, ReportSourceEntry, ReportSourceSummary, StoredReport
from app.report.service import ReportService
from app.report.storage import ReportStorage
from app.session.adapters.sqlite_store import SQLiteSessionStore
from app.session.service import SessionService
from tests.test_helpers import create_test_user_id, patch_test_auth_service, register_and_login

client = TestClient(app)


def build_report(*, report_id: str = "report-export-demo", session_id: str = "session-export") -> StoredReport:
    timestamp = datetime.now(timezone.utc)
    return StoredReport(
        report_id=report_id,
        session_id=session_id,
        report_type="carbon_summary",
        title="碳核算结果说明",
        content=(
            "# 碳核算结果说明\n\n"
            "这是报告正文。\n\n"
            "## 排放明细\n\n"
            "| 排放源 | 活动量 | 排放量 |\n"
            "| --- | ---: | ---: |\n"
            "| 外购电力 | 1000 kWh | 570 kgCO2e |\n\n"
            "- 建议核对电力发票\n"
            "- 复核天然气口径\n"
        ),
        output_format="markdown",
        citations=[
            ReportCitation(
                source_type="carbon_factor",
                title="外购电力因子",
                source="CarbonRag 因子库",
                source_url=None,
                snippet="外购电力排放因子 0.5703 kgCO2e/kWh",
                factor_id="factor-electricity",
            )
        ],
        source_summary=ReportSourceSummary(
            public_policy_count=0,
            private_sample_count=0,
            carbon_factor_count=1,
            total_citation_count=1,
        ),
        sources=[
            ReportSourceEntry(
                source_type="carbon_result",
                source_ref="trace-carbon-001",
                label="2026 核算结果",
                order_index=0,
            )
        ],
        trace_id="report-trace-export",
        created_at=timestamp,
        updated_at=timestamp,
    )


def build_export_service(tmp_path, owner_user_id: str):
    db_path = tmp_path / "carbonrag.sqlite3"
    store = SQLiteSessionStore(db_path)
    store.create_session(
        session_id="session-export",
        owner_user_id=owner_user_id,
        title="导出测试会话",
        created_at=datetime.now(timezone.utc).isoformat(),
    )
    report_storage = ReportStorage(store=store)
    report_service = ReportService(session_service=SessionService(store=store), storage=report_storage)
    created = report_storage.create_report(owner_user_id=owner_user_id, report=build_report())
    export_service = ReportExportService(
        report_service=report_service,
        storage=ReportExportStorage(store=store),
        output_root=tmp_path / "report_files",
    )
    return store, report_service, export_service, created


def test_markdown_ir_parse_headings_and_table() -> None:
    report = build_report()
    detail = report.model_dump()
    from app.report.schemas import ReportDetail

    ir = report_to_ir(ReportDetail.model_validate(detail))

    assert ir.title == "碳核算结果说明"
    assert any(section.heading == "排放明细" for section in ir.sections)
    tables = [block.table for section in ir.sections for block in section.blocks if block.type == "table"]
    assert tables
    assert tables[0].columns == ["排放源", "活动量", "排放量"]
    assert tables[0].rows[0][0] == "外购电力"


def test_docx_export_creates_file_with_table(tmp_path) -> None:
    from app.report.schemas import ReportDetail

    ir = report_to_ir(ReportDetail.model_validate(build_report().model_dump()))
    output_path = tmp_path / "report.docx"
    result = DocxReportExporter().export(ir, output_path)

    assert result.file_size_bytes > 0
    document = Document(output_path)
    assert document.tables
    assert document.tables[1].cell(1, 0).text == "外购电力"


def test_pdf_export_creates_file(tmp_path) -> None:
    from app.report.schemas import ReportDetail

    ir = report_to_ir(ReportDetail.model_validate(build_report().model_dump()))
    output_path = tmp_path / "report.pdf"
    result = PdfReportExporter().export(ir, output_path)

    assert output_path.exists()
    assert result.content_type == "application/pdf"
    assert result.file_size_bytes > 0


def test_export_reuses_existing_file_and_force_regenerates(tmp_path) -> None:
    db_path = tmp_path / "carbonrag.sqlite3"
    owner_user_id = create_test_user_id(db_path, prefix="export-owner")
    _, _, export_service, report = build_export_service(tmp_path, owner_user_id)

    first = export_service.create_exports(
        owner_user_id=owner_user_id,
        report_id=report.report_id,
        payload=CreateReportExportRequest(formats=["docx"]),
    )
    second = export_service.create_exports(
        owner_user_id=owner_user_id,
        report_id=report.report_id,
        payload=CreateReportExportRequest(formats=["docx"]),
    )
    regenerated = export_service.create_exports(
        owner_user_id=owner_user_id,
        report_id=report.report_id,
        payload=CreateReportExportRequest(formats=["docx"], force_regenerate=True),
    )

    assert first.files[0].file_id == second.files[0].file_id
    assert regenerated.files[0].file_id != first.files[0].file_id


def test_export_routes_download_and_reject_other_user(monkeypatch, tmp_path) -> None:
    db_path = tmp_path / "carbonrag.sqlite3"
    patch_test_auth_service(monkeypatch, db_path=db_path)

    user_a = register_and_login(client, prefix="export-a")
    _, _, export_service, report = build_export_service(tmp_path, user_a["user_id"])
    monkeypatch.setattr("app.api.v1.endpoints.report_exports.get_report_export_service", lambda: export_service)

    response = client.post(
        f"/api/v1/reports/{report.report_id}/exports",
        json={"formats": ["docx", "pdf"], "template_id": "default"},
    )
    assert response.status_code == 200, response.text
    files = response.json()["files"]
    assert len(files) == 2

    download_response = client.get(files[0]["download_url"])
    assert download_response.status_code == 200
    assert download_response.headers["content-type"].startswith(files[0]["content_type"])
    assert "attachment" in download_response.headers["content-disposition"]

    register_and_login(client, prefix="export-b")
    assert client.get(files[0]["download_url"]).status_code == 404
    assert client.get(f"/api/v1/reports/{report.report_id}/exports").status_code == 404
