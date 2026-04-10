from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
import xlwt

from app.knowledge.extractor import extract_text_from_source


def _build_pdf(path: Path, text: str) -> None:
    stream = f"BT /F1 18 Tf 72 720 Td ({text}) Tj T* (carbon retrofit) Tj ET".encode("utf-8")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]

    content = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(body)
        content.extend(b"\nendobj\n")

    xref_start = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_start}\n%%EOF"
        ).encode("ascii")
    )
    path.write_bytes(bytes(content))


def _build_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_paragraph(text)
    document.add_paragraph("节能改造建议")
    document.save(str(path))


def _build_xlsx(path: Path, text: str) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "能耗月报"
    sheet.append(["指标", "数值"])
    sheet.append([text, 1200])
    workbook.save(str(path))


def _build_xls(path: Path, text: str) -> None:
    workbook = xlwt.Workbook()
    sheet = workbook.add_sheet("能耗月报")
    sheet.write(0, 0, "指标")
    sheet.write(0, 1, "数值")
    sheet.write(1, 0, text)
    sheet.write(1, 1, 880)
    workbook.save(str(path))


def test_extract_text_from_supported_knowledge_files(tmp_path) -> None:
    txt_path = tmp_path / "sample.txt"
    md_path = tmp_path / "sample.md"
    csv_path = tmp_path / "sample.csv"
    docx_path = tmp_path / "sample.docx"
    xlsx_path = tmp_path / "sample.xlsx"
    xls_path = tmp_path / "sample.xls"
    pdf_path = tmp_path / "sample.pdf"

    txt_path.write_text("双碳目标\n企业节能改造", encoding="utf-8")
    md_path.write_text("# 双碳目标\n\n企业节能改造", encoding="utf-8")
    csv_path.write_text("月份,电量\n2026-01,1200\n2026-02,980\n", encoding="utf-8")
    _build_docx(docx_path, "企业能源审计摘要")
    _build_xlsx(xlsx_path, "电量")
    _build_xls(xls_path, "天然气")
    _build_pdf(pdf_path, "dual carbon target")

    assert "双碳目标" in extract_text_from_source(txt_path)
    assert "企业节能改造" in extract_text_from_source(md_path)
    csv_text = extract_text_from_source(csv_path)
    assert "第 1 行" in csv_text
    assert "电量=1200" in csv_text
    assert "企业能源审计摘要" in extract_text_from_source(docx_path)
    xlsx_text = extract_text_from_source(xlsx_path)
    assert "工作表" in xlsx_text
    assert "数值=1200" in xlsx_text
    xls_text = extract_text_from_source(xls_path)
    assert "天然气" in xls_text
    assert "880" in xls_text
    assert "dual carbon target" in extract_text_from_source(pdf_path)


def test_extract_text_from_old_doc_raises_clear_error(tmp_path) -> None:
    doc_path = tmp_path / "legacy.doc"
    doc_path.write_text("旧版 DOC 文件", encoding="utf-8")

    with pytest.raises(ValueError, match="DOC 旧版二进制文档暂不支持"):
        extract_text_from_source(doc_path)
